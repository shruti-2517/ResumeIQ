"""Unit tests for Stage 2 ATS Readability & Layout Hazard Engine."""

from app.services.parser import check_ats_readability, check_docx_ats, check_pdf_ats


def test_unsupported_file_type():
    result = check_ats_readability(b"test data", "txt")
    assert result["ats_score"] == 0
    assert result["readability_level"] == "Low"
    assert "UNSUPPORTED_TYPE" in result["layout_hazards"]


def test_empty_docx_ats():
    from docx import Document
    import io

    doc = Document()
    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph("Built FastAPI microservices with PostgreSQL and MongoDB.")
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Bachelor of Science in Computer Science.")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, FastAPI, SQL, Docker, React.")

    stream = io.BytesIO()
    doc.save(stream)
    docx_bytes = stream.getvalue()

    result = check_docx_ats(docx_bytes)
    assert result["ats_score"] >= 80
    assert result["readability_level"] == "High"
    assert "Experience" in result["found_sections"]
    assert "Education" in result["found_sections"]
    assert "Skills" in result["found_sections"]


def test_docx_with_table_warning():
    from docx import Document
    import io

    doc = Document()
    doc.add_heading("Experience", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Role"
    table.cell(0, 1).text = "Software Engineer"
    doc.add_heading("Education", level=1)
    doc.add_heading("Skills", level=1)

    stream = io.BytesIO()
    doc.save(stream)
    docx_bytes = stream.getvalue()

    result = check_docx_ats(docx_bytes)
    assert "TABLE_STRUCTURES" in result["layout_hazards"]
    assert any(w["category"] == "Formatting" for w in result["warnings"])
