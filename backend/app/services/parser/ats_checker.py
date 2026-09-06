"""ATS Readability & Layout Hazard Checker for PDF and DOCX resumes."""

import io
import logging
import re
from typing import Any

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)

STANDARD_SECTION_HEADERS = {
    "summary": ["summary", "profile", "professional summary", "about me", "overview"],
    "experience": ["experience", "work experience", "professional experience", "employment", "work history"],
    "education": ["education", "academic background", "qualifications"],
    "skills": ["skills", "technical skills", "core competencies", "technologies"],
    "projects": ["projects", "personal projects", "key projects"],
}


def _detect_pdf_columns(page: Any) -> bool:
    """Detect if a PDF page uses a multi-column layout based on word bounding boxes."""
    try:
        words = page.extract_words()
        if not words or len(words) < 20:
            return False

        page_width = page.width or 600
        midpoint = page_width / 2

        left_count = 0
        right_count = 0

        # Check distribution of words strictly on the left vs right sides
        for word in words:
            x0 = word.get("x0", 0)
            x1 = word.get("x1", 0)
            if x1 < midpoint - 20:
                left_count += 1
            elif x0 > midpoint + 20:
                right_count += 1

        # If significant word clusters exist independently on left & right halves
        total = len(words)
        if (left_count / total > 0.25) and (right_count / total > 0.25):
            return True
    except Exception as exc:
        logger.warning("Error detecting PDF columns: %s", exc)
    return False


def check_pdf_ats(file_bytes: bytes) -> dict[str, Any]:
    """Inspect PDF layout structure for ATS parsing compatibility."""
    score = 100
    warnings: list[dict[str, str]] = []
    hazards: list[str] = []
    found_headers: list[str] = []

    try:
        full_text = ""
        has_multi_column = False
        has_tables = False

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                return {
                    "ats_score": 0,
                    "readability_level": "Low",
                    "warnings": [{"severity": "CRITICAL", "message": "PDF contains no readable pages."}],
                    "found_sections": [],
                    "layout_hazards": ["EMPTY_DOCUMENT"],
                }

            for idx, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                full_text += page_text + "\n"

                # Check columns
                if _detect_pdf_columns(page):
                    has_multi_column = True

                # Check tables
                tables = page.extract_tables()
                if tables and len(tables) > 0:
                    has_tables = True

        if has_multi_column:
            score -= 25
            hazards.append("MULTI_COLUMN_LAYOUT")
            warnings.append(
                {
                    "severity": "HIGH",
                    "category": "Layout",
                    "message": "Multi-column layout detected. Legacy ATS parsers often scramble text across columns.",
                }
            )

        if has_tables:
            score -= 15
            hazards.append("TABLE_STRUCTURES")
            warnings.append(
                {
                    "severity": "MEDIUM",
                    "category": "Formatting",
                    "message": "Tables detected. Information inside table cells may be omitted by some ATS systems.",
                }
            )

        # Check section headers
        text_lower = full_text.lower()
        for category, variations in STANDARD_SECTION_HEADERS.items():
            if any(var in text_lower for var in variations):
                found_headers.append(category.capitalize())
            else:
                if category in ("experience", "education", "skills"):
                    score -= 10
                    warnings.append(
                        {
                            "severity": "HIGH",
                            "category": "Headers",
                            "message": f"Standard section header for '{category.capitalize()}' was not recognized.",
                        }
                    )

        # Check total length
        char_count = len(full_text.strip())
        if char_count < 250:
            score -= 20
            warnings.append(
                {
                    "severity": "HIGH",
                    "category": "Content",
                    "message": "Document text is suspiciously short (< 250 characters). It may be image-based.",
                }
            )

        final_score = max(0, min(100, score))
        level = "High" if final_score >= 80 else ("Medium" if final_score >= 60 else "Low")

        return {
            "ats_score": final_score,
            "readability_level": level,
            "warnings": warnings,
            "found_sections": found_headers,
            "layout_hazards": hazards,
        }

    except Exception as exc:
        logger.error("PDF ATS check failed: %s", exc, exc_info=True)
        return {
            "ats_score": 50,
            "readability_level": "Low",
            "warnings": [{"severity": "HIGH", "category": "Error", "message": f"ATS inspection error: {exc}"}],
            "found_sections": [],
            "layout_hazards": ["PARSE_ERROR"],
        }


def check_docx_ats(file_bytes: bytes) -> dict[str, Any]:
    """Inspect DOCX layout structure for ATS parsing compatibility."""
    score = 100
    warnings: list[dict[str, str]] = []
    hazards: list[str] = []
    found_headers: list[str] = []

    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # Check tables
        if doc.tables and len(doc.tables) > 0:
            score -= 15
            hazards.append("TABLE_STRUCTURES")
            warnings.append(
                {
                    "severity": "MEDIUM",
                    "category": "Formatting",
                    "message": f"{len(doc.tables)} table(s) found. Ensure key content is not trapped in complex tables.",
                }
            )

        # Check headers
        text_lower = full_text.lower()
        for category, variations in STANDARD_SECTION_HEADERS.items():
            if any(var in text_lower for var in variations):
                found_headers.append(category.capitalize())
            else:
                if category in ("experience", "education", "skills"):
                    score -= 10
                    warnings.append(
                        {
                            "severity": "HIGH",
                            "category": "Headers",
                            "message": f"Standard section header for '{category.capitalize()}' was not recognized.",
                        }
                    )

        final_score = max(0, min(100, score))
        level = "High" if final_score >= 80 else ("Medium" if final_score >= 60 else "Low")

        return {
            "ats_score": final_score,
            "readability_level": level,
            "warnings": warnings,
            "found_sections": found_headers,
            "layout_hazards": hazards,
        }
    except Exception as exc:
        logger.error("DOCX ATS check failed: %s", exc, exc_info=True)
        return {
            "ats_score": 50,
            "readability_level": "Low",
            "warnings": [{"severity": "HIGH", "category": "Error", "message": f"ATS inspection error: {exc}"}],
            "found_sections": [],
            "layout_hazards": ["PARSE_ERROR"],
        }


def check_ats_readability(file_bytes: bytes, file_type: str) -> dict[str, Any]:
    """Master router function for ATS Readability Check."""
    if file_type.lower() == "pdf":
        return check_pdf_ats(file_bytes)
    elif file_type.lower() == "docx":
        return check_docx_ats(file_bytes)
    else:
        return {
            "ats_score": 0,
            "readability_level": "Low",
            "warnings": [{"severity": "CRITICAL", "message": f"Unsupported file type: {file_type}"}],
            "found_sections": [],
            "layout_hazards": ["UNSUPPORTED_TYPE"],
        }
